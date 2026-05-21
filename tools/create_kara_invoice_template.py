from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    Image as PdfImage,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "templates"
ASSET_DIR = OUT_DIR / "assets"
DOCX_PATH = OUT_DIR / "kara-rechnungsvorlage.docx"
PDF_PATH = OUT_DIR / "kara-rechnungsvorlage.pdf"
LOGO_DARK = ASSET_DIR / "kara-logo-dark.png"
LOGO_LIGHT = ASSET_DIR / "kara-logo-light.png"

INK = "12110F"
MUTED = "746D63"
LINE = "DED8CF"
SURFACE = "FFFAF2"
SURFACE_DEEP = "EEE6D9"
CHARCOAL = "181715"
CHAMPAGNE = "B89B6F"


def font_path(name):
    path = Path("C:/Windows/Fonts") / name
    return str(path) if path.exists() else None


def create_logo(path, color):
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGBA", (680, 260), (255, 255, 255, 0))
    draw = ImageDraw.Draw(image)
    fill = tuple(int(color[i : i + 2], 16) for i in (0, 2, 4)) + (255,)
    sails = [
        [(96, 20), (60, 72), (76, 126), (165, 151), (132, 91), (148, 28)],
        [(169, 24), (132, 75), (149, 128), (239, 151), (205, 94), (222, 32)],
        [(242, 32), (205, 82), (222, 131), (310, 151), (278, 98), (294, 42)],
    ]
    for points in sails:
        draw.polygon(points, fill=fill)

    font_file = font_path("georgiab.ttf") or font_path("georgia.ttf")
    title_font = ImageFont.truetype(font_file, 82) if font_file else ImageFont.load_default()
    draw.text((50, 146), "Kara", fill=fill, font=title_font)
    image.save(path)


def tighten_doc_spacing(doc):
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=LINE, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn("w:" + side))
        if node is None:
            node = OxmlElement("w:" + side)
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, size=10, bold=False, color=INK, align=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_width(cell, width):
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def add_label(paragraph, text):
    run = paragraph.add_run(text.upper())
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(CHAMPAGNE)


def add_docx():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.05)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    styles["Normal"].font.color.rgb = RGBColor.from_string(INK)
    styles["Title"].font.name = "Georgia"
    styles["Title"].font.size = Pt(34)
    styles["Title"].font.color.rgb = RGBColor.from_string(INK)

    header = doc.add_table(rows=1, cols=2)
    header.autofit = False
    header_cells = header.rows[0].cells
    for cell in header_cells:
        set_cell_shading(cell, CHARCOAL)
        set_cell_border(cell, CHARCOAL)
        set_cell_margins(cell, top=170, bottom=170, start=220, end=220)
    set_width(header_cells[0], Inches(3.9))
    set_width(header_cells[1], Inches(3.15))
    p = header_cells[0].paragraphs[0]
    p.add_run().add_picture(str(LOGO_LIGHT), width=Inches(1.45))
    p2 = header_cells[0].add_paragraph()
    add_label(p2, "Luxury Streetwear / Order Automation")
    p3 = header_cells[1].paragraphs[0]
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p3.add_run("RECHNUNG")
    r.font.name = "Georgia"
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string("FFFFFF")
    p4 = header_cells[1].add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p4.add_run("{{invoice_number}}")
    r.font.name = "Aptos"
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(CHAMPAGNE)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)

    address = doc.add_table(rows=1, cols=2)
    address.autofit = False
    for cell in address.rows[0].cells:
        set_cell_border(cell, "FFFFFF", "0")
        set_cell_margins(cell, start=0, end=240, top=40, bottom=40)
    set_width(address.rows[0].cells[0], Inches(3.95))
    set_width(address.rows[0].cells[1], Inches(3.1))

    left = address.rows[0].cells[0]
    add_label(left.paragraphs[0], "Rechnung an")
    for text, bold in [
        ("{{customer_name}}", True),
        ("{{billing_street}}", False),
        ("{{billing_postal_code}} {{billing_city}}", False),
        ("{{billing_country}}", False),
        ("{{customer_email}}", False),
    ]:
        p = left.add_paragraph()
        r = p.add_run(text)
        r.font.name = "Aptos"
        r.font.size = Pt(10.5)
        r.font.bold = bold

    right = address.rows[0].cells[1]
    right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_label(right.paragraphs[0], "Aussteller")
    for text, bold in [
        ("Kara Studio", True),
        ("Luxury Streetwear Demonstrator", False),
        ("Tiefenbronner Str. 65", False),
        ("75175 Pforzheim, Deutschland", False),
        ("billing@kara.example", False),
    ]:
        p = right.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(text)
        r.font.name = "Aptos"
        r.font.size = Pt(9.6)
        r.font.bold = bold
        r.font.color.rgb = RGBColor.from_string(MUTED if not bold else INK)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)

    facts = doc.add_table(rows=2, cols=4)
    facts.autofit = False
    labels = ["Rechnungsdatum", "Bestellnummer", "Zahlungsstatus", "Versand"]
    values = ["{{invoice_date}}", "{{order_number}}", "Bezahlt", "Standard 48h"]
    for row in facts.rows:
        for cell in row.cells:
            set_cell_border(cell)
            set_cell_margins(cell, top=115, bottom=115, start=120, end=120)
            set_width(cell, Inches(1.75))
    for i, label in enumerate(labels):
        set_cell_text(facts.cell(0, i), label.upper(), size=7.6, bold=True, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(facts.cell(0, i), SURFACE_DEEP)
        set_cell_text(facts.cell(1, i), values[i], size=9.5, bold=True, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER)

    intro = doc.add_paragraph()
    intro.paragraph_format.space_before = Pt(8)
    intro.paragraph_format.space_after = Pt(6)
    intro.add_run("Vielen Dank fuer deinen Einkauf bei Kara. ").bold = True
    intro.add_run(
        "Diese Rechnung wurde automatisch nach erfolgreichem Checkout erzeugt und kann von Make mit den untenstehenden Platzhaltern befüllt werden."
    )

    table = doc.add_table(rows=4, cols=5)
    table.autofit = False
    widths = [Inches(2.15), Inches(2.15), Inches(0.75), Inches(1.0), Inches(1.0)]
    headers = ["Produkt", "Beschreibung", "Menge", "Einzelpreis", "Summe"]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_border(cell)
            set_cell_margins(cell, top=95, bottom=95, start=115, end=115)
            set_width(cell, widths[idx])
    for idx, text in enumerate(headers):
        set_cell_shading(table.cell(0, idx), CHARCOAL)
        set_cell_text(table.cell(0, idx), text, size=8, bold=True, color="FFFFFF")
    rows = [
        ["{{item_name_1}}", "{{item_desc_1}}", "{{qty_1}}", "{{unit_net_1}}", "{{line_net_1}}"],
        ["{{item_name_2}}", "{{item_desc_2}}", "{{qty_2}}", "{{unit_net_2}}", "{{line_net_2}}"],
        ["{{items_note}}", "Weitere Positionen können über Make dupliziert werden.", "", "", ""],
    ]
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, text in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.RIGHT if c_idx >= 2 else None
            set_cell_text(table.cell(r_idx, c_idx), text, size=9, color=INK if r_idx < 3 else MUTED, align=align)
            if r_idx == 3:
                set_cell_shading(table.cell(r_idx, c_idx), SURFACE)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)
    totals = doc.add_table(rows=4, cols=2)
    totals.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for row in totals.rows:
        for cell in row.cells:
            set_cell_border(cell, SURFACE_DEEP)
            set_cell_margins(cell, top=85, bottom=85, start=140, end=140)
    total_rows = [
        ("Zwischensumme netto", "{{subtotal_net}}"),
        ("Versand netto", "{{shipping_net}}"),
        ("Umsatzsteuer 19%", "{{vat_amount}}"),
        ("Gesamtbetrag", "{{total_gross}}"),
    ]
    for idx, (label, value) in enumerate(total_rows):
        fill = CHAMPAGNE if idx == 3 else SURFACE_DEEP
        color = "FFFFFF" if idx == 3 else INK
        for cell in totals.rows[idx].cells:
            set_cell_shading(cell, fill)
        set_cell_text(totals.cell(idx, 0), label, size=9.2, bold=idx == 3, color=color)
        set_cell_text(totals.cell(idx, 1), value, size=10, bold=True, color=color, align=WD_ALIGN_PARAGRAPH.RIGHT)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)
    info = doc.add_table(rows=1, cols=3)
    info.autofit = False
    info_data = [
        ("Zahlung", "Der Betrag wurde im Checkout bezahlt. Es ist keine weitere Überweisung erforderlich."),
        ("Rücksendung", "Rücksendungen sind innerhalb der kommunizierten Frist über den Kara Support anzumelden."),
        ("Automatisierung", "Supabase speichert Bestellung und Positionen; Make versendet Rechnung und Bestätigung."),
    ]
    for idx, cell in enumerate(info.rows[0].cells):
        set_cell_shading(cell, SURFACE)
        set_cell_border(cell)
        set_cell_margins(cell, top=100, bottom=100, start=130, end=130)
        set_width(cell, Inches(2.35))
        add_label(cell.paragraphs[0], info_data[idx][0])
        p = cell.add_paragraph()
        r = p.add_run(info_data[idx][1])
        r.font.name = "Aptos"
        r.font.size = Pt(8.7)
        r.font.color.rgb = RGBColor.from_string(MUTED)

    legal = doc.add_paragraph()
    legal.paragraph_format.space_before = Pt(6)
    legal.paragraph_format.space_after = Pt(2)
    r = legal.add_run(
        "Hinweis: Dies ist eine Vorlage für den IDP-Demonstrator. Alle Platzhalter in doppelten geschweiften Klammern werden durch die Automatisierung ersetzt."
    )
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Kara Studio - Highest Quality since 2016 - International Shipping - Easy Returns")
    r.font.name = "Aptos"
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(MUTED)

    tighten_doc_spacing(doc)
    doc.save(DOCX_PATH)


def p_style(name, size=9, color_hex=INK, leading=None, align=None, font="Helvetica"):
    return ParagraphStyle(
        name,
        fontName=font,
        fontSize=size,
        leading=leading or size * 1.32,
        textColor=colors.HexColor("#" + color_hex),
        alignment=align or 0,
        spaceAfter=0,
    )


def add_pdf():
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        rightMargin=16 * mm,
        leftMargin=16 * mm,
        topMargin=15 * mm,
        bottomMargin=15 * mm,
    )
    story = []
    styles = getSampleStyleSheet()
    styles.add(p_style("Label", 7.2, CHAMPAGNE, 9, font="Helvetica-Bold"))
    styles.add(p_style("BodyKara", 9.3, INK, 12))
    styles.add(p_style("Muted", 8.3, MUTED, 11))
    styles.add(p_style("Right", 9.4, INK, 12, TA_RIGHT))
    styles.add(p_style("InvoiceTitle", 24, "FFFFFF", 28, TA_RIGHT, "Times-Bold"))
    styles.add(p_style("Total", 10, "FFFFFF", 13, TA_RIGHT, "Helvetica-Bold"))

    header = Table(
        [
            [
                PdfImage(str(LOGO_LIGHT), width=47 * mm, height=18 * mm),
                [Paragraph("RECHNUNG", styles["InvoiceTitle"]), Paragraph("{{invoice_number}}", styles["Label"])],
            ]
        ],
        colWidths=[96 * mm, 66 * mm],
        rowHeights=[32 * mm],
    )
    header.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#" + CHARCOAL)),
                ("BOX", (0, 0), (-1, -1), 0, colors.HexColor("#" + CHARCOAL)),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN", (1, 0), (1, 0), "RIGHT"),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
            ]
        )
    )
    story.append(header)
    story.append(Spacer(1, 10 * mm))

    address = Table(
        [
            [
                [
                    Paragraph("RECHNUNG AN", styles["Label"]),
                    Paragraph("<b>{{customer_name}}</b>", styles["BodyKara"]),
                    Paragraph("{{billing_street}}", styles["BodyKara"]),
                    Paragraph("{{billing_postal_code}} {{billing_city}}", styles["BodyKara"]),
                    Paragraph("{{billing_country}}", styles["BodyKara"]),
                    Paragraph("{{customer_email}}", styles["Muted"]),
                ],
                [
                    Paragraph("AUSSTELLER", styles["Label"]),
                    Paragraph("<b>Kara Studio</b>", styles["Right"]),
                    Paragraph("Luxury Streetwear Demonstrator", styles["Right"]),
                    Paragraph("Tiefenbronner Str. 65", styles["Right"]),
                    Paragraph("75175 Pforzheim, Deutschland", styles["Right"]),
                    Paragraph("billing@kara.example", styles["Right"]),
                ],
            ]
        ],
        colWidths=[90 * mm, 72 * mm],
    )
    address.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP")]))
    story.append(address)
    story.append(Spacer(1, 9 * mm))

    facts = Table(
        [
            ["RECHNUNGSDATUM", "BESTELLNUMMER", "ZAHLUNGSSTATUS", "VERSAND"],
            ["{{invoice_date}}", "{{order_number}}", "Bezahlt", "Standard 48h"],
        ],
        colWidths=[40.5 * mm] * 4,
    )
    facts.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#" + SURFACE_DEEP)),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#" + MUTED)),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, 0), 7.2),
                ("FONTNAME", (0, 1), (-1, 1), "Helvetica-Bold"),
                ("FONTSIZE", (0, 1), (-1, 1), 8.8),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#" + LINE)),
                ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#" + LINE)),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ]
        )
    )
    story.append(facts)
    story.append(Spacer(1, 8 * mm))
    story.append(
        Paragraph(
            "<b>Vielen Dank für deinen Einkauf bei Kara.</b> Diese Rechnung wurde automatisch nach erfolgreichem Checkout erzeugt und kann von Make mit den Platzhaltern befüllt werden.",
            styles["BodyKara"],
        )
    )
    story.append(Spacer(1, 6 * mm))

    items = Table(
        [
            ["Produkt", "Beschreibung", "Menge", "Einzelpreis", "Summe"],
            ["{{item_name_1}}", "{{item_desc_1}}", "{{qty_1}}", "{{unit_net_1}}", "{{line_net_1}}"],
            ["{{item_name_2}}", "{{item_desc_2}}", "{{qty_2}}", "{{unit_net_2}}", "{{line_net_2}}"],
            ["{{items_note}}", "Weitere Positionen können über Make dupliziert werden.", "", "", ""],
        ],
        colWidths=[36 * mm, 54 * mm, 17 * mm, 28 * mm, 27 * mm],
    )
    items.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#" + CHARCOAL)),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, 0), 8),
                ("FONTSIZE", (0, 1), (-1, -1), 8.1),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#" + LINE)),
                ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#" + LINE)),
                ("BACKGROUND", (0, 3), (-1, 3), colors.HexColor("#" + SURFACE)),
                ("TEXTCOLOR", (0, 3), (-1, 3), colors.HexColor("#" + MUTED)),
                ("ALIGN", (2, 1), (-1, -1), "RIGHT"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
            ]
        )
    )
    story.append(items)
    story.append(Spacer(1, 7 * mm))

    total = Table(
        [
            ["Zwischensumme netto", "{{subtotal_net}}"],
            ["Versand netto", "{{shipping_net}}"],
            ["Umsatzsteuer 19%", "{{vat_amount}}"],
            ["Gesamtbetrag", "{{total_gross}}"],
        ],
        colWidths=[49 * mm, 34 * mm],
        hAlign="RIGHT",
    )
    total.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 2), colors.HexColor("#" + SURFACE_DEEP)),
                ("BACKGROUND", (0, 3), (-1, 3), colors.HexColor("#" + CHAMPAGNE)),
                ("TEXTCOLOR", (0, 3), (-1, 3), colors.white),
                ("FONTNAME", (0, 3), (-1, 3), "Helvetica-Bold"),
                ("FONTNAME", (1, 0), (1, -1), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("ALIGN", (1, 0), (1, -1), "RIGHT"),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#" + SURFACE_DEEP)),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
                ("LEFTPADDING", (0, 0), (-1, -1), 9),
                ("RIGHTPADDING", (0, 0), (-1, -1), 9),
            ]
        )
    )
    story.append(total)
    story.append(Spacer(1, 8 * mm))

    note_data = [
        [Paragraph("ZAHLUNG", styles["Label"]), Paragraph("RÜCKSENDUNG", styles["Label"]), Paragraph("AUTOMATISIERUNG", styles["Label"])],
        [
            Paragraph("Der Betrag wurde im Checkout bezahlt. Keine weitere Überweisung erforderlich.", styles["Muted"]),
            Paragraph("Rücksendungen werden über den Kara Support angemeldet.", styles["Muted"]),
            Paragraph("Supabase speichert Bestellung und Positionen; Make versendet Rechnung und Bestätigung.", styles["Muted"]),
        ],
    ]
    notes = Table(note_data, colWidths=[54 * mm] * 3)
    notes.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#" + SURFACE)),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#" + LINE)),
                ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#" + LINE)),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    story.append(notes)
    story.append(Spacer(1, 6 * mm))
    story.append(
        Paragraph(
            "Hinweis: Dies ist eine Vorlage für den IDP-Demonstrator. Alle Platzhalter in doppelten geschweiften Klammern werden durch die Automatisierung ersetzt.",
            styles["Muted"],
        )
    )
    story.append(Spacer(1, 4 * mm))
    story.append(Paragraph("Kara Studio - Highest Quality since 2016 - International Shipping - Easy Returns", p_style("Footer", 7.5, MUTED, 9, TA_CENTER)))
    doc.build(story)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    create_logo(LOGO_DARK, INK)
    create_logo(LOGO_LIGHT, "FFFFFF")
    add_docx()
    add_pdf()
    print(DOCX_PATH)
    print(PDF_PATH)


if __name__ == "__main__":
    main()
